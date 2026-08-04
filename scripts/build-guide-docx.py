"""
Build the Word version of the Module & Flow Guide.

The HTML guide in apps/web/public/user-guide.html is the source of truth. This
script parses it and emits a real .docx — not an HTML file renamed, which makes
Word warn the recipient that the format does not match the extension.

Run: python scripts/build-guide-docx.py
"""

from __future__ import annotations

import re
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "apps" / "web" / "public" / "user-guide.html"
OUT = ROOT / "docs" / "Garage-ERP-Guide.docx"

BLUE = RGBColor(0x0F, 0x62, 0xFE)
DARK = RGBColor(0x0F, 0x17, 0x2A)
GREY = RGBColor(0x64, 0x74, 0x8B)
OK = RGBColor(0x19, 0x80, 0x38)
WARN = RGBColor(0xB4, 0x53, 0x09)


class Node:
    """A minimal element tree — enough to walk the guide's structure."""

    def __init__(self, tag: str, attrs: dict[str, str]):
        self.tag = tag
        self.attrs = attrs
        self.children: list["Node | str"] = []

    def text(self) -> str:
        out = []
        for c in self.children:
            out.append(c if isinstance(c, str) else c.text())
        return re.sub(r"\s+", " ", "".join(out)).strip()

    def find_all(self, tag: str) -> list["Node"]:
        found = []
        for c in self.children:
            if isinstance(c, Node):
                if c.tag == tag:
                    found.append(c)
                found.extend(c.find_all(tag))
        return found


class Parser(HTMLParser):
    VOID = {"br", "img", "meta", "link", "hr", "input"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.root = Node("root", {})
        self.stack = [self.root]
        self.skip = 0

    def handle_starttag(self, tag, attrs):
        if tag in ("style", "script", "head"):
            self.skip += 1
            return
        if self.skip or tag in self.VOID:
            return
        node = Node(tag, dict(attrs))
        self.stack[-1].children.append(node)
        self.stack.append(node)

    def handle_endtag(self, tag):
        if tag in ("style", "script", "head"):
            self.skip = max(0, self.skip - 1)
            return
        if self.skip or tag in self.VOID:
            return
        for i in range(len(self.stack) - 1, 0, -1):
            if self.stack[i].tag == tag:
                del self.stack[i:]
                break

    def handle_data(self, data):
        if not self.skip and data.strip():
            self.stack[-1].children.append(data)


def shade(cell, hexcolor: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def status_colour(text: str) -> RGBColor | None:
    t = text.upper()
    if "LIVE" in t:
        return OK
    if "STATIC" in t or "PARTIAL" in t:
        return WARN
    if "NOT BUILT" in t or "PENDING" in t:
        return GREY
    return None


def add_rich(par, node: Node) -> None:
    """Copy a node's inline content, keeping bold, code and status tags."""
    for child in node.children:
        if isinstance(child, str):
            txt = re.sub(r"\s+", " ", child)
            if txt.strip():
                par.add_run(txt)
            continue

        if child.tag == "span" and "tag" in child.attrs.get("class", ""):
            run = par.add_run(f" [{child.text()}] ")
            run.bold = True
            run.font.size = Pt(8)
            colour = status_colour(child.text())
            if colour:
                run.font.color.rgb = colour
            continue

        if child.tag in ("b", "strong"):
            par.add_run(child.text()).bold = True
            continue

        if child.tag in ("em", "i"):
            par.add_run(child.text()).italic = True
            continue

        if child.tag == "code":
            run = par.add_run(child.text())
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        add_rich(par, child)


def build() -> None:
    parser = Parser()
    parser.feed(SRC.read_text(encoding="utf-8"))

    page = None
    for div in parser.root.find_all("div"):
        if "page" in div.attrs.get("class", ""):
            page = div
            break
    if page is None:
        raise SystemExit("could not find the guide body in the HTML")

    doc = Document()

    section = doc.sections[0]
    section.left_margin = section.right_margin = Inches(0.8)
    section.top_margin = section.bottom_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    def walk(node: Node) -> None:
        for child in node.children:
            if isinstance(child, str):
                continue

            tag = child.tag
            cls = child.attrs.get("class", "")

            if tag == "h1":
                h = doc.add_heading(child.text(), level=0)
                for run in h.runs:
                    run.font.color.rgb = DARK

            elif tag == "h2":
                doc.add_page_break() if child.attrs.get("id") == "flows" else None
                h = doc.add_heading(child.text(), level=1)
                for run in h.runs:
                    run.font.color.rgb = BLUE

            elif tag in ("h3", "h4"):
                doc.add_heading(child.text(), level=2 if tag == "h3" else 3)

            elif tag == "p":
                if not child.text():
                    continue
                par = doc.add_paragraph()
                if "sub" in cls:
                    run = par.add_run(child.text())
                    run.font.color.rgb = GREY
                    run.italic = True
                else:
                    add_rich(par, child)

            elif tag in ("div",) and ("lede" in cls or "note" in cls):
                par = doc.add_paragraph()
                par.paragraph_format.left_indent = Inches(0.2)
                add_rich(par, child)
                for run in par.runs:
                    run.font.size = Pt(9.5)

            elif tag == "div" and "kpi" in cls:
                nums = [d for d in child.children if isinstance(d, Node)]
                par = doc.add_paragraph()
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                bits = []
                for d in nums:
                    b = d.find_all("b")
                    s = d.find_all("span")
                    if b and s:
                        bits.append(f"{b[0].text()} {s[0].text()}")
                run = par.add_run("   ·   ".join(bits))
                run.bold = True
                run.font.color.rgb = BLUE

            elif tag == "div" and "chain" in cls:
                steps = [i.text() for i in child.find_all("i")]
                par = doc.add_paragraph()
                par.paragraph_format.left_indent = Inches(0.2)
                run = par.add_run("  →  ".join(steps))
                run.bold = True
                run.font.color.rgb = BLUE
                run.font.size = Pt(9.5)

            elif tag == "div" and "flow" in cls:
                for n, step in enumerate(
                    [s for s in child.children if isinstance(s, Node)], start=1
                ):
                    par = doc.add_paragraph()
                    par.paragraph_format.left_indent = Inches(0.25)
                    par.paragraph_format.space_after = Pt(4)
                    num = par.add_run(f"{n}.  ")
                    num.bold = True
                    num.font.color.rgb = BLUE
                    add_rich(par, step)

            elif tag == "div" and "toc" in cls:
                for li in child.find_all("li"):
                    p = doc.add_paragraph(li.text(), style="List Number")
                    p.paragraph_format.space_after = Pt(0)

            elif tag == "table":
                rows = child.find_all("tr")
                if not rows:
                    continue
                cols = max(
                    len([c for c in r.children if isinstance(c, Node) and c.tag in ("td", "th")])
                    for r in rows
                )
                table = doc.add_table(rows=0, cols=cols)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for r in rows:
                    cells = [c for c in r.children if isinstance(c, Node) and c.tag in ("td", "th")]
                    row = table.add_row()
                    for i, c in enumerate(cells):
                        if i >= cols:
                            break
                        cell = row.cells[i]
                        cell.text = ""
                        par = cell.paragraphs[0]
                        par.paragraph_format.space_after = Pt(2)
                        add_rich(par, c)
                        for run in par.runs:
                            run.font.size = Pt(9)
                            if c.tag == "th":
                                run.bold = True
                        if c.tag == "th":
                            shade(cell, "F1F5F9")
                doc.add_paragraph()

            elif tag in ("ul", "ol"):
                style = "List Bullet" if tag == "ul" else "List Number"
                for li in child.find_all("li"):
                    par = doc.add_paragraph(style=style)
                    par.paragraph_format.space_after = Pt(2)
                    add_rich(par, li)

            elif tag == "footer":
                doc.add_paragraph()
                par = doc.add_paragraph()
                run = par.add_run(child.text())
                run.font.size = Pt(8)
                run.font.color.rgb = GREY

            else:
                walk(child)

    walk(page)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT.relative_to(ROOT)}  ({OUT.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    build()
